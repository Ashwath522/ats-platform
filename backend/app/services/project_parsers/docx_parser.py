"""
backend/app/services/project_parsers/docx_parser.py
──────────────────────────────────────────────────
Extracts text content from Word (.docx / .doc) documents.
Guarantees safe extraction without raising exceptions.
"""
import os
import logging

logger = logging.getLogger("project_parsers.docx")


def parse_docx(file_path: str) -> str:
    """Extract all text from paragraphs and tables in a DOCX file."""
    if not file_path or not os.path.exists(file_path):
        return ""

    try:
        import docx
        doc = docx.Document(file_path)
        parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        # Extract table cells
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)

        return "\n\n".join(parts).strip()
    except Exception as e:
        logger.warning(f"[DOCX_PARSER] python-docx extraction error on {file_path}: {e}")
        # Fallback to text reading
        try:
            with open(file_path, "rb") as f:
                raw = f.read().decode("latin-1", errors="ignore")
            import re
            words = re.findall(r"[A-Za-z0-9_. -]{4,}", raw)
            return " ".join(words[:2000]).strip()
        except Exception as fallback_err:
            logger.error(f"[DOCX_PARSER] Complete fallback failure: {fallback_err}")
            return ""
