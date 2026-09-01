"""
parsers/docx_parser.py
──────────────────────
Word .docx and legacy .doc parsing with fallback string stream extraction.
Pulls every paragraph, table cell, and text segment.
"""

import os
import re


def parse_docx(file_path: str) -> tuple[str, str]:
    """Return (text, method) for a .docx or .doc file."""
    if not file_path or not os.path.exists(file_path):
        return (f"File not found: {file_path}", "missing")

    # 1. Try standard python-docx parser
    try:
        from docx import Document
        doc = Document(file_path)
        parts = []

        # Body paragraphs
        for para in doc.paragraphs:
            if para.text and para.text.strip():
                parts.append(para.text.strip())

        # Table cells
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        text = "\n".join(parts).strip()
        if text:
            return (text, "docx")
    except Exception:
        pass

    # 2. Robust fallback for legacy .doc (binary Word format) or partial documents
    try:
        with open(file_path, "rb") as f:
            raw = f.read()

        # Extract printable ASCII byte sequences
        matches = re.findall(rb"[\x20-\x7E\t\r\n]{4,}", raw)
        lines = []
        for b in matches:
            decoded = b.decode("latin-1", errors="ignore").strip()
            if len(decoded) > 3 and not decoded.startswith(("\\", "{\\", "}}")):
                lines.append(decoded)

        # Filter out common binary noise
        cleaned = [l for l in lines if not any(garbage in l for garbage in ("CompObj", "WordDocument", "SummaryInformation"))]
        text = "\n".join(cleaned[:600]).strip()
        if text:
            return (text, "doc_extracted")
    except Exception as exc:
        return (f"Word document read error: {exc}", "doc_error")

    return (f"Word document: {os.path.basename(file_path)}", "doc_fallback")
